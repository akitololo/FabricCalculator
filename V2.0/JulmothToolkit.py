import os
import sys
import time
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from PyQt4 import QtGui, QtCore

class Window(QtGui.QWidget):
	def __init__(self):
		super(Window, self).__init__()
#	[!] Note that all prices (except blinds and labour) is in metres
		self.FABRIC_PRICE = 1000
		self.LINING_PRICE = 300
		self.BLACKOUT_PRICE = 1300
		self.TAPE_PRICE = 275
		self.SHEERS_PRICE = 900
		self.BLINDS_PRICE = 4500 # Just a reminder: This is per m^2
		self.LABOUR_PRICE = 700 # This one is per Wd (width [according to Dad])
		self.initUI()

	def initUI(self):
		self.setStyleSheet('font-size: 12pt; font-family: Comic Sans MS;')
		self.setGeometry(100, 100, 250, 350)
		self.setWindowTitle("The Fabriccounting Toolkit")
		self.setWindowIcon(QtGui.QIcon('curtain-938541_1280.jpg'))
		self.initUserSpace()

	def initUserSpace(self):
		vbox = QtGui.QVBoxLayout()
		headerLabel = QtGui.QLabel("\t\t\tPlease input the measurements of the below opening in cm\n\t\tPlease enter the SKE & PR numbers with no spaces and in full, e.g.: SKE34003A...\nInvoice number can have leading zeros where necessary. Don't forget to add a description to the document itself!")
		lengthLabel = QtGui.QLabel("Length:")
		widthLabel = QtGui.QLabel("Width:")
		panelLabel = QtGui.QLabel("No. of panels:")
		skeLabel = QtGui.QLabel("SKE Number:")
		prLabel = QtGui.QLabel("PR Number:")
		fboLabel = QtGui.QLabel("FBO Number:")
		invLabel = QtGui.QLabel("Invoice Number:")
		self.lengthInput = QtGui.QLineEdit(self)
		self.widthInput = QtGui.QLineEdit(self)
		self.panelInput = QtGui.QLineEdit(self)
		self.skeInput = QtGui.QLineEdit(self)
		self.prInput = QtGui.QLineEdit(self)
		self.fboInput = QtGui.QLineEdit(self)
		self.invInput = QtGui.QLineEdit(self)
		userInputSpace = QtGui.QFormLayout()
		userInputSpace.addRow(lengthLabel, self.lengthInput)
		userInputSpace.addRow(widthLabel, self.widthInput)
		userInputSpace.addRow(panelLabel, self.panelInput)
		userInputSpace.addRow(skeLabel, self.skeInput)
		userInputSpace.addRow(prLabel, self.prInput)
		userInputSpace.addRow(fboLabel, self.fboInput)
		userInputSpace.addRow(invLabel, self.invInput)
		vbox.addWidget(headerLabel)
		vbox.addLayout(userInputSpace)
		btn = QtGui.QPushButton("Calculate and Generate", self)
		vbox.addWidget(btn)
		btn.clicked.connect(self.retrieveCalculateShow)
		self.popup = None
		self.setLayout(vbox)
		self.show()

	def retrieveCalculateShow(self):
		self.l = float(self.lengthInput.text())
		self.w = float(self.widthInput.text())
		self.p = float(self.panelInput.text())
		Wd = (((self.w)*2) + (10*(self.p))) / 140 # No unit for this particular mini calc
		L = (((self.l) + 28) * Wd) / 100 # This calc ends up in metres (m)
		fabric = lining = blackout = sheers = L
		tape = ((self.w)*2) / 100 # Comes out in m
		labour = Wd
		blinds = (self.w*self.l)/(100**2) # Comes out in m^2

		fabricPrice = fabric * self.FABRIC_PRICE
		liningPrice = lining * self.LINING_PRICE
		blackoutPrice = blackout * self.BLACKOUT_PRICE
		sheersPrice = sheers * self.SHEERS_PRICE
		tapePrice = tape * self.TAPE_PRICE
		labourPrice = labour * self.LABOUR_PRICE
		blindsPrice = blinds * self.BLINDS_PRICE
		totalPrice = fabricPrice + liningPrice + blackoutPrice + sheersPrice + tapePrice + labourPrice + blindsPrice

		fabricCalculatedPriceLabel = QtGui.QLabel(str(int(fabricPrice)))
		liningCalculatedPriceLabel = QtGui.QLabel(str(int(liningPrice)))
		blackoutCalculatedPriceLabel = QtGui.QLabel(str(int(blackoutPrice)))
		sheersCalculatedPriceLabel = QtGui.QLabel(str(int(sheersPrice)))
		tapeCalculatedPriceLabel = QtGui.QLabel(str(int(tapePrice)))
		labourCalculatedPriceLabel = QtGui.QLabel(str(int(labourPrice)))
		blindsCalculatedPriceLabel = QtGui.QLabel(str(int(blindsPrice)))
		totalCalculatedPriceLabel = QtGui.QLabel(str(int(totalPrice)))

##############################################################################################
#							Document generation process begins here 						 #
##############################################################################################

		rootdir = 'C:\\Users\\Austin\\Desktop\\V2.0\\'
		document = Document(os.path.join(rootdir, "default.docx"))
		employer = '\nAMERICAN EMBASSY\t\t\t\t\t\t\t'
		const = 'SKE50016A0101\n'
		pobox = 'P.O.BOX 606-00621\t\t\t\t\t\t\t'
		self.ske_input = str(self.skeInput.text())
		input1 = self.ske_input + '\n'
		loc = 'VILLAGE MARKET\t\t\t\t\t\t\t'
		self.pr_input = str(self.prInput.text())
		input2 = self.pr_input + '\n'
		country = 'NAIROBI KENYA\t\t\t\t\t\t\t'

		day = time.strftime("%d")
		month = time.strftime("%B")
		year = time.strftime("%Y")
		date = day + ' ' + month + ', ' + year + '\n'

		table = document.add_table(rows=5, cols=6)
		table.style = 'TableGrid'
		table_title_start = table.cell(0, 0)
		table_title_end = table.cell(0, 5)
		table_title = table_title_start.merge(table_title_end)
		table_title = table_title.add_paragraph()
		title_format = table_title.paragraph_format
		title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		table_title = table_title.add_run('QUOTATION\n')
		table_title.font.size = Pt(16)

		table_stamp_start = table.cell(1, 0)
		table_stamp_end = table.cell(1, 5)
		table_stamp = table_stamp_start.merge(table_stamp_end)
		table_stamp.text = employer + const + pobox + '\n' + loc + '\n' + country + date

		hdr_cells = table.rows[2].cells
		hdr_cells[0].text = 'ITEM'
		hdr_cells[1].text = 'DESCRIPTION'
		hdr_cells[2].text = 'QTY'
		hdr_cells[3].text = 'UNIT'
		hdr_cells[4].text = 'RATE\n(KSHS)'
		hdr_cells[5].text = 'AMOUNT\n(KSHS)'

		content_cells = table.rows[3].cells
		self.fbo_input = str(self.fboInput.text())
		input3 = '\n' + self.fbo_input + '\n'
		items = '\nFabric\nSheers\nLining\nTape\nBlackout\nVenetian blinds\nLabour\n\nTOTAL\n\nVAT 16%\n\n\n\n\n\n\n'
		content_cells[1].text = input3 + items
		input4 = '\n\n\n' + str(round(fabric, 2)) + '\n'
		input5 = str(round(sheers, 2)) + '\n'
		input6 = str(round(lining, 2)) + '\n'
		input7 = str(round(tape, 2)) + '\n'
		input8 = str(round(blackout, 2)) + '\n'
		input9 = str(round(blinds, 2)) + '\n'
		input10 = str(round(labour, 2)) + '\n'
		content_cells[2].text = input4 + input5 + input6 + input7 + input8 + input9 + input10
		unit1 = '\n\n\nMtrs\n'
		unit2 = unit3 = unit4 = unit5 = 'Mtrs\n'
		unit6 = 'Sqm\n'
		unit7 = 'Wd\n'
		content_cells[3].text = unit1 + unit2 + unit3 + unit4 + unit5 + unit6 + unit7
		rate_fabric = '\n\n\n1,000.00\n'
		rate_sheers = '900.00\n'
		rate_lining = '300.00\n'
		rate_tape = '275.00\n'
		rate_blackout = '850.00\n'
		rate_blinds = '4,500.00\n'
		rate_labcurt = '700.00\n'
		content_cells[4].text = rate_fabric + rate_sheers + rate_lining + rate_tape + rate_blackout + rate_blinds + rate_labcurt
		calc1 = '\n\n\n' + str(round(fabricPrice, 2)) + '\n'
		calc2 = str(round(sheersPrice, 2)) + '\n'
		calc3 = str(round(liningPrice, 2)) + '\n'
		calc4 = str(round(tapePrice, 2)) + '\n'
		calc5 = str(round(blackoutPrice, 2)) + '\n'
		calc6 = str(round(blindsPrice, 2)) + '\n'
		calc7 = str(round(labourPrice, 2)) + '\n' # INQUIRY!! Be sure to ask specifically about this price and whether it requires a separate input field (different labour types)
		total = '\n' + str(round(totalPrice, 2)) + '\n'
		vatCalc = 0.16 *  float(totalPrice)
		vat = '\n' + str(round(vatCalc, 2)) + '\n'
		content_cells[5].text = calc1 + calc2 + calc3 + calc4 + calc5 + calc6 + calc7 + total + vat
		table.cell(4, 1).text = '\nTOTAL'
		total = '\n' + str(round(totalPrice, 2))
		table.cell(4, 5).text = total

		document.save( 'quotation_gen.docx' )

		document1 = Document(os.path.join(rootdir, "default.docx"))

		self.inv_input = str(self.invInput.text())
		inv_num = self.inv_input
		table = document1.add_table(rows=5, cols=6)
		table.style = 'TableGrid'
		table_title_start = table.cell(0, 0)
		table_title_end = table.cell(0, 5)
		table_title = table_title_start.merge(table_title_end)
		table_title = table_title.add_paragraph()
		title_format = table_title.paragraph_format
		title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		table_title = table_title.add_run('INVOICE-' + inv_num + '\n')
		table_title.font.size = Pt(16)

		table_stamp_start = table.cell(1, 0)
		table_stamp_end = table.cell(1, 5)
		table_stamp = table_stamp_start.merge(table_stamp_end)
		table_stamp.text = employer + const + pobox + input1 + loc + input2 + country + date

		hdr_cells = table.rows[2].cells
		hdr_cells[0].text = 'ITEM'
		hdr_cells[1].text = 'DESCRIPTION'
		hdr_cells[2].text = 'QTY'
		hdr_cells[3].text = 'UNIT'
		hdr_cells[4].text = 'RATE\n(KSHS)'
		hdr_cells[5].text = 'AMOUNT\n(KSHS)'

		content_cells = table.rows[3].cells
		txt = '\n(Type here)\n\n\n\n\n\n\n'
		values = 'TOTAL\n\n\nVAT 16%\n\n\n\n\n\n\n'
		content_cells[1].text = input3 + txt + values
		total = '\n\n\n\n\n\n\n' + str(round(totalPrice, 2)) + '\n\n'
		total2 = '\n' + str(round(totalPrice, 2)) + '\n(make bold and underline)\n\n'
		vat = '\n' + str(round(vatCalc, 2))
		content_cells[5].text = total + total2 + vat

		table.cell(4, 1).text = '\nTOTAL'
		total = '\n' + str(round(totalPrice, 2))
		table.cell(4, 5).text = total

		document1.save( 'invoice_gen.docx' )

##############################################################################################
#							Pop-up handling begins here 									 #
##############################################################################################

		popup = QtGui.QDialog()
		popup.setStyleSheet('font-size: 12pt; font-family: Comic Sans MS;')
		popup.setGeometry(150,150, 200, 230)
		popup.setWindowIcon(QtGui.QIcon('curtain-938541_1280.jpg'))
		popup.setWindowTitle("Your Answers")

		answerSpace = QtGui.QFormLayout()
		fabricSizeLabel = QtGui.QLabel("Fabric (m):")
		liningSizeLabel = QtGui.QLabel("Lining (m):")
		blackoutSizeLabel = QtGui.QLabel("Blackout (m):")
		tapeSizeLabel = QtGui.QLabel("Tape (m):")
		sheersSizeLabel = QtGui.QLabel("Sheers (m):")
		blindsSizeLabel = QtGui.QLabel("Blinds (m^2):")
		labourSizeLabel = QtGui.QLabel("Labour (Wd):")
		fabricPriceLabel = QtGui.QLabel("Fabric price total (Ksh):")
		liningPriceLabel = QtGui.QLabel("Lining price total (Ksh):")
		blackoutPriceLabel = QtGui.QLabel("Blackout price total (Ksh):")
		tapePriceLabel = QtGui.QLabel("Tape price total (Ksh):")
		sheersPriceLabel = QtGui.QLabel("Sheers price total (Ksh):")
		blindsPriceLabel = QtGui.QLabel("Blinds price total (Ksh):")
		labourPriceLabel = QtGui.QLabel("Labour price total (Ksh):")
		totalPriceLabel = QtGui.QLabel("Total (Ksh):")
		line = QtGui.QLabel("   ______________________")
		success = QtGui.QLabel("Documents processed successfully")
		answerSpace.addRow(fabricSizeLabel, QtGui.QLabel(str(round(fabric, 3))))
		answerSpace.addRow(liningSizeLabel, QtGui.QLabel(str(round(lining, 3))))
		answerSpace.addRow(blackoutSizeLabel, QtGui.QLabel(str(round(blackout, 3))))
		answerSpace.addRow(tapeSizeLabel, QtGui.QLabel(str(round(tape, 3))))
		answerSpace.addRow(sheersSizeLabel, QtGui.QLabel(str(round(sheers, 3))))
		answerSpace.addRow(blindsSizeLabel, QtGui.QLabel(str(round(blinds, 3))))
		answerSpace.addRow(labourSizeLabel, QtGui.QLabel(str(round(labour, 3))))
		answerSpace.addRow(line) #<- Seems to be acting as a gap currently. Looks fine. Doesn't appear to cause any issues that aren't aesthetic
		answerSpace.addRow(fabricPriceLabel, fabricCalculatedPriceLabel)
		answerSpace.addRow(liningPriceLabel, liningCalculatedPriceLabel)
		answerSpace.addRow(blackoutPriceLabel, blackoutCalculatedPriceLabel)
		answerSpace.addRow(tapePriceLabel, tapeCalculatedPriceLabel)
		answerSpace.addRow(sheersPriceLabel, sheersCalculatedPriceLabel)
		answerSpace.addRow(blindsPriceLabel, blindsCalculatedPriceLabel)
		answerSpace.addRow(labourPriceLabel, labourCalculatedPriceLabel)
		answerSpace.addRow(totalPriceLabel, totalCalculatedPriceLabel)
		answerSpace.addRow(line)
		answerSpace.addRow(success)


		popup.setLayout(answerSpace)

		popup.exec_()


def main():
	app = QtGui.QApplication(sys.argv)
	GUI = Window()
	sys.exit(app.exec_())

if __name__ == '__main__':
	main()