from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
import time

document = Document()
employer = '\nAMERICAN EMBASSY\t\t\t\t\t\t\t'
const = 'SKE50016A0101\n'
pobox = 'P.O.BOX 606-00621\t\t\t\t\t\t\t'
input1 = 'SKE NUM\n'
loc = 'VILLAGE MARKET\t\t\t\t\t\t\t'
input2 = 'PR NUM\n'
country = 'NAIROBI KENYA\t\t\t\t\t\t\t'

day = time.strftime("%d")
month = time.strftime("%B")
year = time.strftime("%Y")
date = day + ' ' + month + ', ' + year + '\n'

table = document.add_table(rows=5, cols=6)
table.style = 'TableGrid'
table_title_start = table.cell(0, 0)
table_title_end = table.cell(0, 5)
########## Possible Shortcut ##########
# row = table.rows[0]
# a, b = row.cell[:6]
# a.merge(b)
#######################################
table_title = table_title_start.merge(table_title_end)
table_title = table_title.add_paragraph()
title_format = table_title.paragraph_format
title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
table_title = table_title.add_run('QUOTATION\n')
table_title.font.size = Pt(16)

table_stamp_start = table.cell(1, 0)
table_stamp_end = table.cell(1, 5)
########## Possible Shortcut ##########
# row = table.rows[1]
# a, b = row.cell[:6]
# a.merge(b)
#######################################
table_stamp = table_stamp_start.merge(table_stamp_end)
table_stamp.text = employer + const + pobox + input1 + loc + input2 + country + date

hdr_cells = table.rows[2].cells
hdr_cells[0].text = 'ITEM'
hdr_cells[1].text = 'DESCRIPTION'
hdr_cells[2].text = 'QTY'
hdr_cells[3].text = 'UNIT'
hdr_cells[4].text = 'RATE\n(KSHS)'
hdr_cells[5].text = 'AMOUNT\n(KSHS)'

content_cells = table.rows[3].cells
input3 = '\nFBO NUM\n'
items = '\nFabric\nSheers\nLining\nTape\nBlackout\nVenetian blinds\nLabour for curtains\nLabour for sheers\n\nTOTAL\n\nVAT 16%\n\n\n\n\n\n\n'
content_cells[1].text = input3 + items
input4 = '\n\n\nNUM\n'
input5, input6, input7, input8, input9, input10, input11 = 'NUM\n', 'NUM\n', 'NUM\n', 'NUM\n', 'NUM\n', 'NUM\n', 'NUM\n'
content_cells[2].text = input4 + input5 + input6 + input7 + input8 + input9 + input10 + input11
unit1 = '\n\n\nMtrs\n'
unit2, unit3, unit4, unit5 = 'Mtrs\n', 'Mtrs\n', 'Mtrs\n', 'Mtrs\n'
unit6 = 'Sqm\n'
unit7, unit8 = 'Wd\n', 'Wd\n'
content_cells[3].text = unit1 + unit2 + unit3 + unit4 + unit5 + unit6 + unit7 + unit8
rate_fabric = '\n\n\n1,000.00\n'
rate_sheers = '900.00\n'
rate_lining = '300.00\n'
rate_tape = '275.00\n'
rate_blackout = '850.00\n'
rate_blinds = '4,500.00\n'
rate_labcurt, rate_labshrs = '700.00\n', '700.00\n'
content_cells[4].text = rate_fabric + rate_sheers + rate_lining + rate_tape + rate_blackout + rate_blinds + rate_labcurt + rate_labshrs
calc1 = '\n\n\nNUM\n'
calc2, calc3, calc4, calc5, calc6, calc7, calc8 = 'NUM\n', 'NUM\n', 'NUM\n', 'NUM\n', 'NUM\n', 'NUM\n', 'NUM\n'
total = '\nSUM\n'
vat = '\nTAX\n'
content_cells[5].text = calc1 + calc2 + calc3 + calc4 + calc5 + calc6 + calc7 + calc8 + total + vat
table.cell(4, 1).text = 'TOTAL'
total = 'SUM'
table.cell(4, 5).text = total

document.save( 'quotation_gen.docx' )


document1 = Document()

inv_num = 'NUM'
table = document1.add_table(rows=5, cols=6)
table.style = 'TableGrid'
table_title_start = table.cell(0, 0)
table_title_end = table.cell(0, 5)
table_title = table_title_start.merge(table_title_end)
table_title = table_title.add_paragraph()
title_format = table_title.paragraph_format
title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
table_title = table_title.add_run('INVOICE-' + inv_num + '\n')
table_title.font.size = Pt(16)

table_stamp_start = table.cell(1, 0)
table_stamp_end = table.cell(1, 5)
table_stamp = table_stamp_start.merge(table_stamp_end)
table_stamp.text = employer + const + pobox + input1 + loc + input2 + country + date

hdr_cells = table.rows[2].cells
hdr_cells[0].text = 'ITEM'
hdr_cells[1].text = 'DESCRIPTION'
hdr_cells[2].text = 'QTY'
hdr_cells[3].text = 'UNIT'
hdr_cells[4].text = 'RATE\n(KSHS)'
hdr_cells[5].text = 'AMOUNT\n(KSHS)'

content_cells = table.rows[3].cells
txt = '\n(Type here)\n\n\n\n\n\n\n'
values = 'TOTAL\n\n\nVAT 16%\n\n\n\n\n\n\n'
content_cells[1].text = input3 + txt + values

total = '\n\n\n\n\n\n\nSUM\n\n'
total2 = '\nSUM_BOLD\n\n'
tax = '\nTAX'
content_cells[5].text = total + total2 + tax

table.cell(4, 1).text = 'TOTAL'
total = 'SUM'
table.cell(4, 5).text = total

document1.save( 'invoice_gen.docx' )
